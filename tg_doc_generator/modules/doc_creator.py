from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cv_template(cv_data, filename):
    # Create a new Document
    doc = Document()
    
    # Add Title
    title = doc.add_heading('Curriculum Vitae', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Personal Details Section
    doc.add_heading('Personal Details', level=2)
    for key, value in cv_data['personal_details'].items():
        p = doc.add_paragraph()
        p.add_run(f'{key}: ').bold = True
        p.add_run(value)

    # Add Professional Experience Section
    doc.add_heading('Professional Experience', level=2)
    for job in cv_data['professional_experience']:
        doc.add_heading(job['title'], level=3)
        doc.add_paragraph(job['company'])
        doc.add_paragraph(job['dates'])
        doc.add_paragraph(job['description'])

    # Add Education Section
    doc.add_heading('Education', level=2)
    for edu in cv_data['education']:
        doc.add_heading(edu['degree'], level=3)
        doc.add_paragraph(edu['institution'])
        doc.add_paragraph(edu['dates'])
        doc.add_paragraph(edu['description'])

    # Add Skills Section
    doc.add_heading('Skills', level=2)
    doc.add_paragraph(', '.join(cv_data['skills']))

    # Save the document
    doc.save(filename)